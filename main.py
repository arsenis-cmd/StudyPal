# nlp_service/main.py - StudyPal NLP Microservice
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional
import PyPDF2
import docx
import spacy
import re
from collections import Counter
import random
from datetime import datetime
import io

app = FastAPI(title="StudyPal NLP API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load spaCy model for NLP
nlp = spacy.load("en_core_web_sm")

class DocumentText(BaseModel):
    text: str
    title: Optional[str] = None

class FlashcardRequest(BaseModel):
    text: str
    count: int = 10
    difficulty: str = "medium"

class QuizRequest(BaseModel):
    text: str
    question_count: int = 10
    quiz_type: str = "multiple_choice"
    difficulty: str = "medium"

class SummarizeRequest(BaseModel):
    text: str
    max_sentences: int = 5

@app.get("/")
def read_root():
    return {"status": "StudyPal NLP API is running", "version": "1.0.0"}

@app.post("/parse-document")
async def parse_document(file: UploadFile = File(...)):
    """
    Extract text from uploaded document (PDF, DOCX, TXT)
    """
    try:
        content = await file.read()
        file_extension = file.filename.split(".")[-1].lower()
        
        extracted_text = ""
        metadata = {
            "filename": file.filename,
            "file_type": file_extension,
            "file_size": len(content)
        }
        
        if file_extension == "pdf":
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            metadata["page_count"] = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n"
                
        elif file_extension == "docx":
            doc = docx.Document(io.BytesIO(content))
            metadata["page_count"] = len(doc.paragraphs)
            
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
                
        elif file_extension == "txt":
            extracted_text = content.decode('utf-8')
            metadata["page_count"] = 1
            
        else:
            raise HTTPException(400, "Unsupported file type. Use PDF, DOCX, or TXT")
        
        # Clean text
        extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
        
        # Calculate metadata
        words = extracted_text.split()
        metadata["word_count"] = len(words)
        metadata["character_count"] = len(extracted_text)
        metadata["estimated_read_time"] = round(len(words) / 200)  # avg 200 wpm
        
        return {
            "text": extracted_text,
            "metadata": metadata
        }
        
    except Exception as e:
        raise HTTPException(500, f"Error parsing document: {str(e)}")

@app.post("/analyze")
def analyze_content(request: DocumentText):
    """
    Analyze document content and extract key information
    """
    try:
        doc = nlp(request.text[:1000000])  # Limit for performance
        
        # Extract named entities
        entities = {}
        for ent in doc.ents:
            if ent.label_ not in entities:
                entities[ent.label_] = []
            if ent.text not in entities[ent.label_]:
                entities[ent.label_].append(ent.text)
        
        # Extract key noun phrases
        noun_phrases = []
        for chunk in doc.noun_chunks:
            if len(chunk.text.split()) >= 2:
                noun_phrases.append(chunk.text)
        
        # Get most common noun phrases (potential concepts)
        noun_phrase_counts = Counter(noun_phrases)
        key_concepts = [phrase for phrase, count in noun_phrase_counts.most_common(20)]
        
        # Identify topics using keyword extraction
        topics = extract_topics(request.text)
        
        # Estimate difficulty
        difficulty = estimate_difficulty(doc)
        
        # Generate summary
        summary = generate_summary(request.text)
        
        return {
            "summary": summary,
            "key_concepts": key_concepts[:15],
            "entities": entities,
            "topics": topics,
            "difficulty_level": difficulty,
            "word_count": len(request.text.split()),
            "estimated_study_time": max(10, len(request.text.split()) // 150)  # minutes
        }
        
    except Exception as e:
        raise HTTPException(500, f"Error analyzing content: {str(e)}")

@app.post("/generate-flashcards")
def generate_flashcards(request: FlashcardRequest):
    """
    Generate flashcards from text using NLP
    """
    try:
        doc = nlp(request.text[:500000])
        flashcards = []
        
        # Strategy 1: Definition-based flashcards
        sentences = [sent.text.strip() for sent in doc.sents]
        
        for sentence in sentences:
            # Look for definition patterns
            if " is " in sentence or " are " in sentence or " means " in sentence:
                parts = re.split(r'\s+(?:is|are|means)\s+', sentence, maxsplit=1)
                if len(parts) == 2:
                    term = parts[0].strip()
                    definition = parts[1].strip()
                    
                    if 5 < len(term.split()) < 10 and 10 < len(definition.split()) < 50:
                        flashcards.append({
                            "front": f"What is {term}?",
                            "back": definition,
                            "context": sentence,
                            "difficulty": assign_difficulty(definition)
                        })
        
        # Strategy 2: Question-answer pairs from important sentences
        important_sentences = get_important_sentences(sentences)
        
        for sent in important_sentences[:request.count]:
            # Extract subject
            sent_doc = nlp(sent)
            subjects = [token.text for token in sent_doc if token.dep_ == "nsubj"]
            
            if subjects:
                subject = subjects[0]
                flashcards.append({
                    "front": f"Explain {subject} in the context of this material",
                    "back": sent,
                    "context": sent,
                    "difficulty": assign_difficulty(sent)
                })
        
        # Strategy 3: Key concept flashcards
        key_terms = extract_key_terms(doc)
        
        for term in key_terms[:request.count // 3]:
            # Find sentences containing the term
            context_sents = [s for s in sentences if term.lower() in s.lower()]
            if context_sents:
                flashcards.append({
                    "front": f"Define or explain: {term}",
                    "back": context_sents[0],
                    "context": " ".join(context_sents[:2]),
                    "difficulty": "medium"
                })
        
        # Limit to requested count
        flashcards = flashcards[:request.count]
        
        # Add IDs
        for i, card in enumerate(flashcards):
            card["id"] = i + 1
        
        return {
            "flashcards": flashcards,
            "total_generated": len(flashcards)
        }
        
    except Exception as e:
        raise HTTPException(500, f"Error generating flashcards: {str(e)}")

@app.post("/generate-quiz")
def generate_quiz(request: QuizRequest):
    """
    Generate quiz questions from text
    """
    try:
        doc = nlp(request.text[:500000])
        questions = []
        
        sentences = [sent.text.strip() for sent in doc.sents if len(sent.text.split()) > 8]
        important_sentences = get_important_sentences(sentences)
        
        # Generate multiple choice questions
        for i, sentence in enumerate(important_sentences[:request.question_count]):
            sent_doc = nlp(sentence)
            
            # Find entities or important nouns to blank out
            entities = [ent.text for ent in sent_doc.ents]
            nouns = [token.text for token in sent_doc if token.pos_ == "NOUN"]
            
            blanks = entities if entities else nouns
            
            if blanks:
                correct_answer = random.choice(blanks)
                
                # Create question by removing the answer
                question_text = sentence.replace(correct_answer, "______")
                
                # Generate distractors (wrong answers)
                all_terms = extract_key_terms(doc)
                distractors = [term for term in all_terms if term != correct_answer]
                random.shuffle(distractors)
                distractors = distractors[:3]
                
                # Create options
                options = [correct_answer] + distractors
                random.shuffle(options)
                
                questions.append({
                    "question_id": i + 1,
                    "question_text": question_text,
                    "question_type": "multiple_choice",
                    "options": options,
                    "correct_answer": correct_answer,
                    "explanation": f"In the text: {sentence}",
                    "difficulty": assign_difficulty(sentence),
                    "points": 10
                })
        
        # Add true/false questions
        for i, sentence in enumerate(sentences[request.question_count:request.question_count + 5]):
            # Randomly make it true or false
            is_true = random.choice([True, False])
            
            if not is_true:
                # Modify sentence to make it false
                sent_doc = nlp(sentence)
                nouns = [token.text for token in sent_doc if token.pos_ == "NOUN"]
                if nouns:
                    wrong_noun = random.choice(extract_key_terms(doc))
                    sentence = sentence.replace(random.choice(nouns), wrong_noun, 1)
            
            questions.append({
                "question_id": len(questions) + 1,
                "question_text": sentence,
                "question_type": "true_false",
                "options": ["True", "False"],
                "correct_answer": "True" if is_true else "False",
                "explanation": "Based on the source material",
                "difficulty": "easy",
                "points": 5
            })
        
        return {
            "questions": questions[:request.question_count],
            "total_questions": min(len(questions), request.question_count),
            "total_points": sum(q["points"] for q in questions[:request.question_count])
        }
        
    except Exception as e:
        raise HTTPException(500, f"Error generating quiz: {str(e)}")

@app.post("/summarize")
def summarize_text(request: SummarizeRequest):
    """
    Generate a summary of the text
    """
    try:
        sentences = re.split(r'[.!?]+', request.text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        # Score sentences by importance
        scored_sentences = get_important_sentences(sentences)
        
        # Return top N sentences
        summary = ". ".join(scored_sentences[:request.max_sentences]) + "."
        
        return {
            "summary": summary,
            "original_length": len(request.text.split()),
            "summary_length": len(summary.split()),
            "compression_ratio": f"{(len(summary.split()) / len(request.text.split()) * 100):.1f}%"
        }
        
    except Exception as e:
        raise HTTPException(500, f"Error summarizing: {str(e)}")

# Helper Functions

def extract_topics(text: str) -> List[str]:
    """Extract main topics from text"""
    doc = nlp(text[:100000])
    
    # Get most frequent nouns and proper nouns
    nouns = [token.lemma_ for token in doc if token.pos_ in ["NOUN", "PROPN"]]
    noun_counts = Counter(nouns)
    
    topics = [noun for noun, count in noun_counts.most_common(10) if count > 2]
    return topics

def estimate_difficulty(doc) -> str:
    """Estimate text difficulty based on various metrics"""
    # Average sentence length
    sentences = list(doc.sents)
    avg_sentence_length = sum(len(sent) for sent in sentences) / len(sentences) if sentences else 0
    
    # Vocabulary complexity (ratio of unique words)
    words = [token.text for token in doc if token.is_alpha]
    vocab_diversity = len(set(words)) / len(words) if words else 0
    
    # Named entities density
    entity_density = len(doc.ents) / len(doc) if len(doc) > 0 else 0
    
    # Scoring
    difficulty_score = (
        (avg_sentence_length / 20) * 0.4 +
        vocab_diversity * 0.4 +
        entity_density * 20 * 0.2
    )
    
    if difficulty_score < 0.4:
        return "easy"
    elif difficulty_score < 0.7:
        return "medium"
    else:
        return "hard"

def generate_summary(text: str, max_sentences: int = 5) -> str:
    """Generate extractive summary"""
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    
    if len(sentences) <= max_sentences:
        return text
    
    scored = get_important_sentences(sentences)
    return ". ".join(scored[:max_sentences]) + "."

def get_important_sentences(sentences: List[str]) -> List[str]:
    """Score and rank sentences by importance"""
    sentence_scores = []
    
    for sentence in sentences:
        score = 0
        words = sentence.lower().split()
        
        # Longer sentences get higher scores
        score += min(len(words) / 10, 2)
        
        # Sentences with numbers/data
        if any(char.isdigit() for char in sentence):
            score += 1
        
        # Sentences with key phrases
        key_phrases = ["important", "significant", "key", "main", "primary", "essential"]
        if any(phrase in sentence.lower() for phrase in key_phrases):
            score += 2
        
        # Avoid questions as summaries
        if sentence.strip().endswith("?"):
            score -= 1
        
        sentence_scores.append((sentence, score))
    
    # Sort by score
    sentence_scores.sort(key=lambda x: x[1], reverse=True)
    return [sent for sent, score in sentence_scores]

def extract_key_terms(doc) -> List[str]:
    """Extract important terms from document"""
    # Nouns and proper nouns
    terms = [token.lemma_ for token in doc if token.pos_ in ["NOUN", "PROPN"] and len(token.text) > 3]
    
    # Named entities
    entities = [ent.text for ent in doc.ents]
    
    all_terms = terms + entities
    term_counts = Counter(all_terms)
    
    return [term for term, count in term_counts.most_common(30) if count > 1]

def assign_difficulty(text: str) -> str:
    """Assign difficulty level to text"""
    words = text.split()
    avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
    
    if avg_word_length < 5:
        return "easy"
    elif avg_word_length < 7:
        return "medium"
    else:
        return "hard"

@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "nlp_model": "en_core_web_sm"
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
